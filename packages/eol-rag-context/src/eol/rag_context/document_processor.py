"""Comprehensive document processing system for multi-format content extraction and
chunking.

This module provides intelligent document processing capabilities that can handle various
file formats including code files, documentation, PDFs, Word documents, and structured
data. It implements advanced chunking strategies optimized for RAG systems, preserving
document structure and semantic boundaries while creating appropriately-sized chunks.

Key Features:
    - Multi-format support: Code, Markdown, PDF, DOCX, JSON/YAML, plain text
    - Intelligent chunking with semantic boundary preservation
    - AST-based code parsing using Tree-sitter for structure-aware chunking
    - Header-based Markdown organization with hierarchical structure
    - Content-aware PDF processing with paragraph detection
    - Configurable chunk sizes with overlap for optimal retrieval
    - Comprehensive metadata extraction for enhanced search capabilities
    - MIME type detection and file size validation

Supported Formats:
    - Code: Python, JavaScript/TypeScript, Rust, Go, Java, C/C++, C#
    - Documentation: Markdown, reStructuredText, plain text
    - Documents: PDF, Microsoft Word (.docx)
    - Structured: JSON, YAML
    - Automatic MIME type detection for unknown extensions

Chunking Strategies:
    - Semantic: Respects paragraph boundaries and document structure
    - Structural: Uses headers, functions, classes as natural boundaries
    - Token-based: Fixed-size chunks with configurable overlap
    - Format-specific: Optimized strategies for each document type

Example:
    Basic document processing workflow:

    >>> from eol.rag_context.document_processor import DocumentProcessor
    >>> from eol.rag_context.config import DocumentConfig, ChunkingConfig
    >>> from pathlib import Path
    >>>
    >>> # Configure processing
    >>> doc_config = DocumentConfig(
    ...     max_file_size_mb=50,
    ...     extract_metadata=True,
    ...     parse_code_structure=True
    ... )
    >>> chunk_config = ChunkingConfig(
    ...     max_chunk_size=1024,
    ...     use_semantic_chunking=True,
    ...     code_chunk_by_function=True
    ... )
    >>>
    >>> # Initialize processor
    >>> processor = DocumentProcessor(doc_config, chunk_config)
    >>>
    >>> # Process different file types
    >>> python_doc = await processor.process_file(Path("main.py"))
    >>> print(f"Code document: {len(python_doc.chunks)} chunks")
    >>>
    >>> markdown_doc = await processor.process_file(Path("README.md"))
    >>> print(f"Markdown: {len(markdown_doc.chunks)} sections")
    >>>
    >>> # Process entire directory
    >>> documents = await processor.process_directory(Path("/project"))
    >>> print(f"Processed {len(documents)} documents")
    >>>
    >>> # Examine processing results
    >>> for doc in documents[:3]:
    ...     print(f"{doc.file_path}: {doc.doc_type} ({len(doc.chunks)} chunks)")
    ...     if doc.language:
    ...         print(f"  Language: {doc.language}")
    ...     print(f"  Metadata: {list(doc.metadata.keys())}")

"""

import json
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import aiofiles
import magic
import markdown
import pypdf
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

# Optional tree-sitter for AST parsing
try:
    from tree_sitter import Language, Parser

    TREE_SITTER_AVAILABLE = True
except ImportError:
    TREE_SITTER_AVAILABLE = False
    Language = None
    Parser = None
import logging

from .config import ChunkingConfig, DocumentConfig

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Comprehensive document representation with content, metadata, and processing
    results.

    Represents a fully processed document containing the original content,
    extracted metadata, intelligent chunks, and format-specific information.
    Used throughout the RAG system for indexing and retrieval operations.

    Attributes:
        file_path: Path to the original source file.
        content: Full document content as extracted text.
        doc_type: Document type classification (markdown, code, pdf, docx, text, structured).
        metadata: Extracted document metadata (title, author, size, format-specific info).
        chunks: List of intelligently created content chunks for indexing.
        language: Programming language for code files (None for non-code documents).

    Example:
        Accessing processed document information:

        >>> doc = await processor.process_file(Path("example.py"))
        >>> print(f"File: {doc.file_path}")
        >>> print(f"Type: {doc.doc_type}, Language: {doc.language}")
        >>> print(f"Content length: {len(doc.content)} chars")
        >>> print(f"Chunks: {len(doc.chunks)}")
        >>>
        >>> # Examine metadata
        >>> metadata = doc.metadata
        >>> print(f"Lines: {metadata.get('lines', 'N/A')}")
        >>> print(f"Format: {metadata.get('format', 'unknown')}")
        >>>
        >>> # Analyze chunks
        >>> for i, chunk in enumerate(doc.chunks[:3]):
        ...     print(f"Chunk {i}: {chunk['type']} ({chunk['tokens']} tokens)")
        ...     print(f"  Content preview: {chunk['content'][:50]}...")

    """

    file_path: Path
    content: str
    doc_type: str  # markdown, code, pdf, docx, text, json
    metadata: dict[str, Any] = field(default_factory=dict)
    chunks: list[dict[str, Any]] = field(default_factory=list)
    language: str | None = None  # For code files


class DocumentProcessor:
    """Advanced multi-format document processor with intelligent chunking for RAG
    systems.

    Provides comprehensive document processing capabilities that handle various file
    formats and implement sophisticated chunking strategies optimized for retrieval
    and semantic search. Integrates with Tree-sitter for AST-based code parsing
    and uses format-specific heuristics for optimal content extraction.

    Processing Pipeline:
    1. File type detection using MIME types and extensions
    2. Format-specific content extraction with metadata
    3. Intelligent chunking based on document structure
    4. Chunk metadata enrichment for enhanced search

    Key Capabilities:
        - Multi-format processing with automatic type detection
        - Structure-aware chunking preserving semantic boundaries
        - AST-based code parsing for function/class level chunks
        - Metadata extraction for enhanced search and filtering
        - Configurable chunk sizes with overlap strategies
        - Error handling and graceful degradation

    Attributes:
        doc_config: Document processing configuration settings.
        chunk_config: Chunking strategy and size configuration.
        mime: Magic library instance for MIME type detection.
        parsers: Dictionary of Tree-sitter parsers for code languages.

    Example:
        Complete document processing setup:

        >>> from eol.rag_context.config import DocumentConfig, ChunkingConfig
        >>>
        >>> # Configure for code-heavy projects
        >>> doc_config = DocumentConfig(
        ...     file_patterns=["*.py", "*.js", "*.md", "*.rst"],
        ...     max_file_size_mb=25,
        ...     parse_code_structure=True,
        ...     extract_metadata=True
        ... )
        >>>
        >>> chunk_config = ChunkingConfig(
        ...     max_chunk_size=800,
        ...     chunk_overlap=100,
        ...     use_semantic_chunking=True,
        ...     code_chunk_by_function=True,
        ...     markdown_split_headers=True
        ... )
        >>>
        >>> processor = DocumentProcessor(doc_config, chunk_config)
        >>>
        >>> # Process with format-specific optimization
        >>> results = []
        >>> for file_path in source_files:
        ...     doc = await processor.process_file(file_path)
        ...     if doc:
        ...         results.append(doc)
        ...         print(f"Processed {doc.doc_type}: {len(doc.chunks)} chunks")
        >>>
        >>> print(f"Total: {len(results)} documents processed")

    """

    def __init__(self, doc_config: DocumentConfig, chunk_config: ChunkingConfig):
        """Initialize document processor with configuration and parsing capabilities.

        Args:
            doc_config: Document processing configuration including file patterns,
                size limits, and processing options.
            chunk_config: Chunking configuration including size limits, overlap
                settings, and strategy preferences.

        """
        self.doc_config = doc_config
        self.chunk_config = chunk_config
        self.mime = magic.Magic(mime=True)
        self.parsers = self._init_code_parsers()

    def _init_code_parsers(self) -> dict[str, Parser]:
        """Initialize tree-sitter parsers for code."""
        parsers = {}

        if not TREE_SITTER_AVAILABLE:
            return parsers

        # Map file extensions to languages
        lang_map = {
            ".py": "python",
            ".js": "javascript",
            ".jsx": "javascript",
            ".ts": "typescript",
            ".tsx": "typescript",
            ".rs": "rust",
            ".go": "go",
            ".java": "java",
        }

        # Load available parsers
        try:
            import tree_sitter_go
            import tree_sitter_java
            import tree_sitter_javascript
            import tree_sitter_python
            import tree_sitter_rust
            import tree_sitter_typescript

            languages = {
                "python": Language(tree_sitter_python.language),
                "javascript": Language(tree_sitter_javascript.language),
                "typescript": Language(tree_sitter_typescript.language_typescript),
                "rust": Language(tree_sitter_rust.language),
                "go": Language(tree_sitter_go.language),
                "java": Language(tree_sitter_java.language),
            }

            for ext, lang in lang_map.items():
                if lang in languages:
                    parser = Parser(languages[lang])
                    parsers[ext] = parser
        except ImportError as e:
            logger.warning(f"Some tree-sitter languages not available: {e}")

        return parsers

    async def process_file(self, file_path: Path) -> ProcessedDocument | None:
        """Process a single file using format-specific extraction and chunking
        strategies.

        Automatically detects file type using MIME detection and file extensions,
        then routes to the appropriate format-specific processor. Handles various
        document types with optimized processing strategies for each format.

        Processing Flow:
        1. Validates file existence and size limits
        2. Detects file type using MIME types and extensions
        3. Routes to format-specific processor
        4. Extracts content and metadata
        5. Applies intelligent chunking strategies
        6. Returns structured document representation

        Args:
            file_path: Path to the file to process. Must exist and be readable.

        Returns:
            ProcessedDocument containing extracted content, metadata, and chunks.
            None if file cannot be processed (too large, unsupported format, errors).

        Raises:
            IOError: If file cannot be read due to permissions or corruption.

        Example:
            Process different document types:

            >>> from pathlib import Path
            >>>
            >>> # Process Python code with AST chunking
            >>> python_doc = await processor.process_file(Path("main.py"))
            >>> if python_doc:
            ...     print(f"Language: {python_doc.language}")
            ...     func_chunks = [c for c in python_doc.chunks if 'function' in c['type']]
            ...     print(f"Functions found: {len(func_chunks)}")
            >>>
            >>> # Process Markdown with header-based chunking
            >>> md_doc = await processor.process_file(Path("README.md"))
            >>> if md_doc:
            ...     headers = md_doc.metadata.get('headers', [])
            ...     print(f"Document structure: {len(headers)} headers")
            >>>
            >>> # Process PDF with page-based extraction
            >>> pdf_doc = await processor.process_file(Path("manual.pdf"))
            >>> if pdf_doc:
            ...     print(f"Pages: {pdf_doc.metadata.get('pages', 0)}")
            ...     print(f"Paragraphs: {len(pdf_doc.chunks)}")

        Note:
            Files exceeding max_file_size_mb configuration are automatically
            skipped with a warning log message.

        """
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.doc_config.max_file_size_mb:
            logger.warning(f"File too large ({file_size_mb:.2f}MB): {file_path}")
            return None

        # Detect file type
        mime_type = self.mime.from_file(str(file_path))
        suffix = file_path.suffix.lower()

        # Route to appropriate processor
        if suffix == ".md":
            return await self._process_markdown(file_path)
        elif suffix == ".pdf":
            return await self._process_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return await self._process_docx(file_path)
        elif suffix in [".json", ".yaml", ".yml"]:
            return await self._process_structured(file_path)
        elif suffix in [".xml", ".rss", ".atom", ".svg", ".xhtml"]:
            return await self._process_xml(file_path)
        elif suffix in [
            ".py",
            ".js",
            ".jsx",
            ".ts",
            ".tsx",
            ".rs",
            ".go",
            ".java",
            ".cpp",
            ".c",
            ".cs",
        ]:
            # Process as code even if tree-sitter parser not available
            return await self._process_code(file_path)
        elif suffix in self.parsers:
            return await self._process_code(file_path)
        elif "text" in mime_type:
            return await self._process_text(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_path} ({mime_type})")
            return None

    async def _process_markdown(self, file_path: Path) -> ProcessedDocument:
        """Process Markdown files with header-based structure preservation.

        Processes Markdown files by parsing to HTML, extracting header structure,
        and creating chunks that respect document hierarchy. Can chunk by headers
        or use semantic text chunking based on configuration.

        Features:
            - HTML parsing for structure extraction
            - Header hierarchy analysis
            - Header-based chunking (when enabled)
            - Metadata extraction including document structure

        Args:
            file_path: Path to Markdown file to process.

        Returns:
            ProcessedDocument with Markdown-specific structure and chunking.

        """
        async with aiofiles.open(file_path, encoding="utf-8") as f:
            content = await f.read()

        # Parse markdown
        html = markdown.markdown(content, extensions=["extra", "codehilite", "toc"])
        soup = BeautifulSoup(html, "html.parser")

        # Extract structure
        doc = ProcessedDocument(
            file_path=file_path,
            content=content,
            doc_type="markdown",
            metadata={
                "format": "markdown",
                "size": len(content),
                "headers": self._extract_headers(soup),
            },
        )

        # Chunk by headers if configured
        if self.chunk_config.markdown_split_headers:
            doc.chunks = self._chunk_markdown_by_headers(content, str(file_path))
        else:
            doc.chunks = self._chunk_text(content, str(file_path))

        return doc

    def _extract_headers(self, soup: BeautifulSoup) -> list[dict[str, Any]]:
        """Extract header structure from HTML."""
        headers = []
        for h in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6"]):
            headers.append(
                {
                    "level": int(h.name[1]),
                    "text": h.get_text().strip(),
                }
            )
        return headers

    def _chunk_markdown_by_headers(
        self, content: str, source_path: str = ""
    ) -> list[dict[str, Any]]:
        """Chunk markdown by header structure."""
        chunks = []
        lines = content.split("\n")
        current_chunk = []
        current_header = None

        for line in lines:
            # Check if line is a header
            header_match = re.match(r"^(#{1,6})\s+(.*)", line)

            if header_match:
                # Save previous chunk if exists
                if current_chunk:
                    chunks.append(
                        self._create_chunk(
                            content="\n".join(current_chunk),
                            chunk_type="section",
                            chunk_index=len(chunks),
                            source=source_path,
                            header=current_header,
                            section=current_header,
                        )
                    )

                # Start new chunk
                current_header = header_match.group(2)
                current_chunk = [line]
            else:
                current_chunk.append(line)

        # Add final chunk
        if current_chunk:
            chunks.append(
                self._create_chunk(
                    content="\n".join(current_chunk),
                    chunk_type="section",
                    chunk_index=len(chunks),
                    source=source_path,
                    header=current_header,
                    section=current_header,
                )
            )

        return chunks

    async def _process_pdf(self, file_path: Path) -> ProcessedDocument:
        """Process PDF files."""
        content_parts = []
        metadata = {"pages": 0, "format": "pdf"}

        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            metadata["pages"] = len(reader.pages)

            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                content_parts.append(text)

                # Extract metadata from first page
                if i == 0 and reader.metadata:
                    metadata.update(
                        {
                            "title": reader.metadata.get("/Title", ""),
                            "author": reader.metadata.get("/Author", ""),
                            "subject": reader.metadata.get("/Subject", ""),
                        }
                    )

        content = "\n\n".join(content_parts)

        doc = ProcessedDocument(
            file_path=file_path, content=content, doc_type="pdf", metadata=metadata
        )

        # Chunk by pages or semantic boundaries
        doc.chunks = self._chunk_pdf_content(content_parts, str(file_path))

        return doc

    def _chunk_pdf_content(self, pages: list[str], source_path: str = "") -> list[dict[str, Any]]:
        """Chunk PDF content intelligently."""
        chunks = []

        for i, page_content in enumerate(pages):
            # Try to detect paragraphs
            paragraphs = re.split(r"\n\s*\n", page_content)

            for para in paragraphs:
                if len(para.strip()) > 50:  # Skip very short paragraphs
                    chunks.append(
                        self._create_chunk(
                            content=para.strip(),
                            chunk_type="paragraph",
                            chunk_index=len(chunks),
                            source=source_path,
                            page=i + 1,
                        )
                    )

        return chunks

    async def _process_docx(self, file_path: Path) -> ProcessedDocument:
        """Process Word documents."""
        doc = DocxDocument(str(file_path))

        content_parts = []
        metadata = {"format": "docx", "properties": {}}

        # Extract core properties
        if doc.core_properties:
            metadata["properties"] = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": (
                    str(doc.core_properties.created) if doc.core_properties.created else ""
                ),
            }

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            content_parts.append("\n".join(table_text))

        content = "\n\n".join(content_parts)

        doc_obj = ProcessedDocument(
            file_path=file_path, content=content, doc_type="docx", metadata=metadata
        )

        doc_obj.chunks = self._chunk_text(content)

        return doc_obj

    async def _process_code(self, file_path: Path) -> ProcessedDocument:
        """Process source code files with AST-based parsing and structure-aware
        chunking.

        Processes source code files using Tree-sitter parsers when available for
        structure-aware chunking by functions, classes, and other code constructs.
        Falls back to line-based chunking if AST parsing is not available.

        Features:
            - Language detection from file extensions
            - AST parsing for semantic boundaries (functions, classes)
            - Line-based fallback chunking with overlap
            - Comprehensive metadata extraction (lines, language, etc.)

        Args:
            file_path: Path to source code file to process.

        Returns:
            ProcessedDocument with language-specific chunking and metadata.

        Note:
            AST parsing requires Tree-sitter languages to be installed.
            Without AST support, falls back to line-based chunking.

        """
        async with aiofiles.open(file_path, encoding="utf-8") as f:
            content = await f.read()

        suffix = file_path.suffix.lower()
        language = self._detect_language(suffix)

        doc = ProcessedDocument(
            file_path=file_path,
            content=content,
            doc_type="code",
            language=language,
            metadata={
                "format": "code",
                "language": language,
                "lines": len(content.splitlines()),
            },
        )

        # Parse with tree-sitter if available
        if suffix in self.parsers and self.chunk_config.code_chunk_by_function:
            doc.chunks = self._chunk_code_by_ast(
                content, self.parsers[suffix], language, str(file_path)
            )
        else:
            doc.chunks = self._chunk_code_by_lines(content, language, str(file_path))

        return doc

    def _detect_language(self, suffix: str) -> str:
        """Detect programming language from file extension."""
        lang_map = {
            ".py": "python",
            ".js": "javascript",
            ".jsx": "javascript",
            ".ts": "typescript",
            ".tsx": "typescript",
            ".rs": "rust",
            ".go": "go",
            ".java": "java",
            ".cpp": "cpp",
            ".c": "c",
            ".cs": "csharp",
            ".rb": "ruby",
            ".php": "php",
            ".swift": "swift",
            ".kt": "kotlin",
        }
        return lang_map.get(suffix, "unknown")

    def _chunk_code_by_ast(
        self, content: str, parser: Parser, language: str, source_path: str = ""
    ) -> list[dict[str, Any]]:
        """Chunk code using AST to preserve function/class boundaries."""
        chunks = []
        tree = parser.parse(bytes(content, "utf8"))

        # Define node types to chunk by
        chunk_node_types = {
            "python": ["function_definition", "class_definition"],
            "javascript": [
                "function_declaration",
                "class_declaration",
                "arrow_function",
            ],
            "typescript": [
                "function_declaration",
                "class_declaration",
                "arrow_function",
            ],
            "rust": ["function_item", "impl_item"],
            "go": ["function_declaration", "method_declaration"],
            "java": ["method_declaration", "class_declaration"],
        }

        target_types = chunk_node_types.get(language, [])

        def extract_nodes(node, depth=0):
            """Recursively extract target nodes."""
            if node.type in target_types:
                start = node.start_byte
                end = node.end_byte
                chunk_content = (
                    content[start:end].decode("utf-8")
                    if isinstance(content[start:end], bytes)
                    else content[start:end]
                )

                chunks.append(
                    self._create_chunk(
                        content=chunk_content,
                        chunk_type=node.type,
                        chunk_index=len(chunks),
                        source=source_path,
                        language=language,
                        start_line=node.start_point[0] + 1,
                        end_line=node.end_point[0] + 1,
                    )
                )

            for child in node.children:
                extract_nodes(child, depth + 1)

        extract_nodes(tree.root_node)

        # If no functions found, fall back to line-based chunking
        if not chunks:
            return self._chunk_code_by_lines(content, language, source_path)

        return chunks

    def _chunk_code_by_lines(
        self, content: str, language: str, source_path: str = ""
    ) -> list[dict[str, Any]]:
        """Chunk code by lines with overlap."""
        chunks = []
        lines = content.splitlines()

        chunk_size = self.chunk_config.code_max_lines
        overlap = min(10, chunk_size // 4)

        for i in range(0, len(lines), chunk_size - overlap):
            chunk_lines = lines[i : i + chunk_size]
            if chunk_lines:
                chunks.append(
                    self._create_chunk(
                        content="\n".join(chunk_lines),
                        chunk_type="lines",
                        chunk_index=len(chunks),
                        source=source_path,
                        language=language,
                        start_line=i + 1,
                        end_line=min(i + chunk_size, len(lines)),
                    )
                )

        return chunks

    async def _process_structured(self, file_path: Path) -> ProcessedDocument:
        """Process JSON/YAML files."""
        async with aiofiles.open(file_path, encoding="utf-8") as f:
            content = await f.read()

        suffix = file_path.suffix.lower()

        # Parse structured data
        if suffix == ".json":
            import json

            data = json.loads(content)
            doc_type = "structured"  # Use "structured" for consistency with tests
        else:
            import yaml

            data = yaml.safe_load(content)
            doc_type = "structured"  # Use "structured" for consistency with tests

        doc = ProcessedDocument(
            file_path=file_path,
            content=content,
            doc_type=doc_type,
            metadata={
                "format": doc_type,
                "keys": list(data.keys()) if isinstance(data, dict) else [],
                "type": type(data).__name__,
            },
        )

        # Chunk by top-level keys or array items
        doc.chunks = self._chunk_structured_data(data, doc_type, str(file_path))

        return doc

    def _chunk_structured_data(
        self, data: Any, format: str, source_path: str = ""
    ) -> list[dict[str, Any]]:
        """Chunk structured data intelligently."""
        chunks = []

        if isinstance(data, dict):
            for key, value in data.items():
                chunk_content = (
                    json.dumps({key: value}, indent=2) if format == "json" else str({key: value})
                )
                chunks.append(
                    self._create_chunk(
                        content=chunk_content,
                        chunk_type="object_field",
                        chunk_index=len(chunks),
                        source=source_path,
                        key=key,
                        format=format,
                    )
                )
        elif isinstance(data, list):
            for i, item in enumerate(data):
                chunk_content = json.dumps(item, indent=2) if format == "json" else str(item)
                chunks.append(
                    self._create_chunk(
                        content=chunk_content,
                        chunk_type="array_item",
                        chunk_index=len(chunks),
                        source=source_path,
                        array_index=i,
                        format=format,
                    )
                )
        else:
            # Single value
            chunks.append(
                self._create_chunk(
                    content=str(data),
                    chunk_type="value",
                    chunk_index=0,
                    source=source_path,
                    format=format,
                )
            )

        return chunks

    async def _process_xml(self, file_path: Path) -> ProcessedDocument:
        """Process XML files with structure preservation.

        Processes XML files by parsing the tree structure, extracting elements
        and attributes, and creating semantic chunks based on XML hierarchy.
        Supports various XML formats including RSS, Atom, SVG, and configuration files.

        Features:
            - XML tree structure parsing and preservation
            - Namespace handling
            - Attribute extraction as metadata
            - XPath tracking for precise location
            - Specialized handling for RSS/Atom feeds
            - Element-based semantic chunking

        Args:
            file_path: Path to the XML file.

        Returns:
            ProcessedDocument with structured XML content and semantic chunks.
        """
        import xml.etree.ElementTree as ET

        async with aiofiles.open(file_path, encoding="utf-8") as f:
            content = await f.read()

        # Parse XML
        try:
            tree = ET.fromstring(content)
        except ET.ParseError as e:
            logger.warning(f"Failed to parse XML {file_path}: {e}")
            # Fall back to text processing
            return await self._process_text(file_path)

        # Extract metadata
        metadata = {
            "format": "xml",
            "root_tag": tree.tag,
            "root_attributes": tree.attrib,
            "element_count": len(tree.findall(".//*")),
        }

        # Extract namespaces
        namespaces = {}
        for elem in tree.iter():
            if elem.tag.startswith("{"):
                ns_end = elem.tag.find("}")
                ns = elem.tag[1:ns_end]
                if ns not in namespaces:
                    namespaces[ns] = ns.split("/")[-1]

        if namespaces:
            metadata["namespaces"] = namespaces

        # Detect special XML types
        suffix = file_path.suffix.lower()
        if suffix in [".rss", ".atom"] or any(
            tag in content[:500] for tag in ["<rss", "<feed", "<channel"]
        ):
            doc_type = "feed"
            chunks = self._chunk_xml_feed(tree, str(file_path))
        elif suffix == ".svg" or "<svg" in content[:100]:
            doc_type = "svg"
            chunks = self._chunk_svg(tree, str(file_path))
        elif any(tag in content[:500] for tag in ["<configuration>", "<config>", "<settings>"]):
            doc_type = "config"
            chunks = self._chunk_xml_config(tree, str(file_path))
        elif tree.tag == "event" or any(
            tag in content[:500] for tag in ["<event>", "<calendar>", "<date>"]
        ):
            doc_type = "event"
            chunks = self._chunk_xml_event(tree, str(file_path))
        else:
            doc_type = "xml"
            chunks = self._chunk_xml_generic(tree, str(file_path))

        # Extract text content for full-text search
        text_content = self._extract_xml_text(tree)

        doc = ProcessedDocument(
            file_path=file_path, content=text_content, doc_type=doc_type, metadata=metadata
        )
        doc.chunks = chunks

        return doc

    def _extract_xml_text(self, element) -> str:
        """Extract all text content from XML element recursively."""
        text_parts = []

        if element.text:
            text_parts.append(element.text.strip())

        for child in element:
            child_text = self._extract_xml_text(child)
            if child_text:
                text_parts.append(child_text)

        if element.tail:
            text_parts.append(element.tail.strip())

        return " ".join(filter(None, text_parts))

    def _extract_temporal_metadata(self, root) -> dict[str, Any]:
        """Extract temporal information from XML document."""
        temporal = {}

        # Common temporal element names
        temporal_tags = [
            "date",
            "time",
            "datetime",
            "calendar",
            "when",
            "schedule",
            "pubdate",
            "published",
            "updated",
            "created",
            "modified",
            "start",
            "end",
            "begin",
            "finish",
            "deadline",
        ]

        for tag in temporal_tags:
            # Search case-insensitively
            elements = root.findall(f".//{tag}")
            if not elements:
                # Try with different case
                elements = root.findall(f".//{tag.capitalize()}")
            if not elements:
                elements = root.findall(f".//{tag.upper()}")

            for elem in elements:
                if elem.text and elem.text.strip():
                    # Store the first found date/time
                    if "date" not in temporal:
                        temporal["date"] = elem.text.strip()
                        temporal["date_source"] = tag
                    # Store all found dates with their tag names
                    temporal[f"{tag}_value"] = elem.text.strip()

        # Also check attributes for temporal information
        for elem in root.iter():
            for attr_name, attr_value in elem.attrib.items():
                if any(t in attr_name.lower() for t in ["date", "time", "when"]):
                    temporal[f"attr_{attr_name}"] = attr_value

        return temporal

    def _chunk_xml_generic(self, root, source_path: str) -> list[dict[str, Any]]:
        """Create chunks from generic XML structure with enhanced temporal context."""
        chunks = []

        # First pass: Extract any temporal information from the document
        temporal_metadata = self._extract_temporal_metadata(root)

        def process_element(elem, xpath="", parent_context="", depth=0, inherited_temporal=None):
            # Build XPath
            tag_name = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            current_xpath = f"{xpath}/{tag_name}" if xpath else tag_name

            # Check if this element contains temporal information
            local_temporal = {}
            if tag_name.lower() in ["date", "time", "datetime", "calendar", "when", "schedule"]:
                if elem.text and elem.text.strip():
                    local_temporal["date"] = elem.text.strip()

            # Merge inherited and local temporal context
            current_temporal = {**(inherited_temporal or {}), **local_temporal}

            # Extract element text
            elem_text = self._extract_xml_text(elem)

            # Determine if element should be a chunk
            if self._should_chunk_xml_element(elem, elem_text):
                chunk_metadata = {
                    "xpath": current_xpath,
                    "tag": tag_name,
                    "attributes": elem.attrib,
                    "depth": depth,
                    "parent_context": parent_context[:200] if parent_context else "",
                    "has_children": len(elem) > 0,
                }

                # Add temporal context to metadata if available
                if current_temporal:
                    chunk_metadata.update(current_temporal)

                # If we have temporal context, add it to the content
                content = elem_text
                if current_temporal and tag_name.lower() not in [
                    "date",
                    "time",
                    "datetime",
                    "calendar",
                ]:
                    temporal_prefix = []
                    if "date" in current_temporal:
                        temporal_prefix.append(f"Date/Time: {current_temporal['date']}")
                    if temporal_prefix:
                        content = "\n".join(temporal_prefix) + "\n\n" + elem_text

                chunks.append(
                    self._create_chunk(
                        content=content,
                        chunk_type="xml_element",
                        chunk_index=len(chunks),
                        source=source_path,
                        **chunk_metadata,
                    )
                )
                # Update parent context for children
                parent_context = elem_text[:200] if elem_text else ""

            # Process children with inherited temporal context
            for child in elem:
                process_element(child, current_xpath, parent_context, depth + 1, current_temporal)

        # Start processing from root with global temporal metadata
        process_element(root, inherited_temporal=temporal_metadata)

        # If no chunks were created, create one for the entire document
        if not chunks:
            chunks.append(
                self._create_chunk(
                    content=self._extract_xml_text(root),
                    chunk_type="xml_document",
                    chunk_index=0,
                    source=source_path,
                    root_tag=root.tag,
                )
            )

        return chunks

    def _should_chunk_xml_element(self, element, text_content: str) -> bool:
        """Determine if an XML element should become a chunk."""
        if not text_content or len(text_content.strip()) < 50:
            return False

        # Common semantic XML elements
        semantic_tags = {
            "paragraph",
            "p",
            "section",
            "article",
            "chapter",
            "description",
            "abstract",
            "summary",
            "content",
            "body",
            "text",
            "note",
            "comment",
            "entry",
            "record",
            "item",
            "row",
            "document",
            "div",
            "message",
            "post",
            "block",
            "field",
        }

        tag_name = element.tag.split("}")[-1].lower() if "}" in element.tag else element.tag.lower()

        if tag_name in semantic_tags:
            return True

        # Chunk if it's a leaf with substantial content
        if len(element) == 0 and len(text_content) > 100:
            return True

        # Chunk if it has few children but substantial content
        if len(element) < 5 and len(text_content) > 200:
            return True

        return False

    def _chunk_xml_feed(self, root, source_path: str) -> list[dict[str, Any]]:
        """Process RSS/Atom feeds."""
        chunks = []

        # Handle RSS - items are under channel
        channel = root.find("channel")
        if channel is not None:
            items = channel.findall("item")
        else:
            # Handle both RSS (without channel) and Atom formats
            items = root.findall(".//item")  # RSS
            if not items:
                items = root.findall(".//{http://www.w3.org/2005/Atom}entry")  # Atom

        for item in items:
            content_parts = []
            metadata = {"type": "feed_item"}

            # Extract common fields
            title = item.find("title")
            if title is None:
                title = item.find("{http://www.w3.org/2005/Atom}title")
            if title is not None and title.text:
                metadata["title"] = title.text
                content_parts.append(f"Title: {title.text}")

            description = item.find("description")
            if description is None:
                description = item.find("{http://www.w3.org/2005/Atom}summary")
            if description is None:
                description = item.find("{http://www.w3.org/2005/Atom}content")
            if description is not None and description.text:
                metadata["description"] = description.text[:500]
                content_parts.append(f"Content: {description.text}")

            link = item.find("link")
            if link is None:
                link = item.find("{http://www.w3.org/2005/Atom}link")
            if link is not None:
                if link.text:
                    metadata["link"] = link.text
                else:
                    metadata["link"] = link.get("href", "")

            pubDate = item.find("pubDate")
            if pubDate is None:
                pubDate = item.find("{http://www.w3.org/2005/Atom}published")
            if pubDate is None:
                pubDate = item.find("{http://www.w3.org/2005/Atom}updated")
            if pubDate is not None and pubDate.text:
                metadata["date"] = pubDate.text

            if content_parts:
                chunks.append(
                    self._create_chunk(
                        content="\n".join(content_parts),
                        chunk_type="feed_item",
                        chunk_index=len(chunks),
                        source=source_path,
                        **metadata,
                    )
                )

        return chunks

    def _chunk_svg(self, root, source_path: str) -> list[dict[str, Any]]:
        """Process SVG files."""
        chunks = []

        # Extract text elements
        text_elements = root.findall(".//{http://www.w3.org/2000/svg}text")
        if not text_elements:
            text_elements = root.findall(".//text")

        text_content = []
        for text_elem in text_elements:
            if text_elem.text:
                text_content.append(text_elem.text)

        # Create a single chunk for SVG
        metadata = {
            "type": "svg",
            "width": root.get("width", "unknown"),
            "height": root.get("height", "unknown"),
            "viewBox": root.get("viewBox", ""),
            "has_text": len(text_content) > 0,
            "text_count": len(text_content),
        }

        content = f"SVG Image: {root.get('id', 'unnamed')}"
        if text_content:
            content += f"\nText elements: {', '.join(text_content)}"

        chunks.append(
            self._create_chunk(
                content=content, chunk_type="svg", chunk_index=0, source=source_path, **metadata
            )
        )

        return chunks

    def _chunk_xml_config(self, root, source_path: str) -> list[dict[str, Any]]:
        """Process configuration XML files."""
        chunks = []

        def extract_config_items(elem, path=""):
            tag_name = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            current_path = f"{path}.{tag_name}" if path else tag_name

            # If leaf node with value
            if not list(elem) and elem.text and elem.text.strip():
                chunks.append(
                    self._create_chunk(
                        content=f"{current_path} = {elem.text.strip()}",
                        chunk_type="config_value",
                        chunk_index=len(chunks),
                        source=source_path,
                        config_path=current_path,
                        value=elem.text.strip(),
                        attributes=elem.attrib,
                    )
                )

            # Process children
            for child in elem:
                extract_config_items(child, current_path)

        extract_config_items(root)

        # If no config items found, treat as generic XML
        if not chunks:
            return self._chunk_xml_generic(root, source_path)

        return chunks

    def _chunk_xml_event(self, root, source_path: str) -> list[dict[str, Any]]:
        """Process event XML files with temporal context preservation.

        For event XMLs, we create a single comprehensive chunk with all information
        since these are typically small, focused documents about single events.
        """
        chunks = []

        # Extract all relevant information
        temporal_context = {}

        # Look for calendar/date information
        calendar_elem = root.find(".//calendar")
        if calendar_elem is not None:
            date_elem = calendar_elem.find("date")
            if date_elem is not None and date_elem.text:
                temporal_context["date"] = date_elem.text
                temporal_context["datetime_raw"] = date_elem.text

            location_elem = calendar_elem.find("location")
            if location_elem is not None and location_elem.text:
                temporal_context["location"] = location_elem.text

        # Also check for direct date elements
        date_elem = root.find(".//date")
        if date_elem is not None and date_elem.text and "date" not in temporal_context:
            temporal_context["date"] = date_elem.text

        # Extract event ID
        event_id = None
        id_elem = root.find("id")
        if id_elem is not None and id_elem.text:
            event_id = id_elem.text
            temporal_context["event_id"] = event_id

        # Extract title
        title = None
        title_elem = root.find("title")
        if title_elem is not None and title_elem.text:
            title = title_elem.text
            temporal_context["title"] = title

        # Extract URL
        url_elem = root.find("url")
        if url_elem is not None and url_elem.text:
            temporal_context["url"] = url_elem.text

        # Extract details
        details = None
        details_elem = root.find("details")
        if details_elem is not None and details_elem.text:
            details = details_elem.text

        # Build comprehensive event content
        content_parts = []

        # Add structured header with key information
        if title:
            content_parts.append(f"Event: {title}")
        if temporal_context.get("date"):
            content_parts.append(f"Date/Time: {temporal_context['date']}")
        if temporal_context.get("location"):
            content_parts.append(f"Location: {temporal_context['location']}")
        if event_id:
            content_parts.append(f"Event ID: {event_id}")

        # Add main details
        if details:
            content_parts.append(f"\nDetails:\n{details}")

        # Add URL if present
        if temporal_context.get("url"):
            content_parts.append(f"\nMore info: {temporal_context['url']}")

        # Create single comprehensive chunk for the entire event
        # This ensures all temporal and event information stays together
        chunks.append(
            self._create_chunk(
                content="\n".join(content_parts),
                chunk_type="event_complete",
                chunk_index=0,
                source=source_path,
                **temporal_context,  # Include all metadata
            )
        )

        return chunks

    async def _process_text(self, file_path: Path) -> ProcessedDocument:
        """Process plain text files."""
        async with aiofiles.open(file_path, encoding="utf-8") as f:
            content = await f.read()

        doc = ProcessedDocument(
            file_path=file_path,
            content=content,
            doc_type="text",
            metadata={
                "format": "text",
                "size": len(content),
                "lines": len(content.splitlines()),
            },
        )

        doc.chunks = self._chunk_text(content, str(file_path))

        return doc

    def _create_chunk(
        self, content: str, chunk_type: str = "text", chunk_index: int = 0, **metadata
    ) -> dict[str, Any]:
        """Create a chunk with consistent metadata structure."""
        return {
            "content": content,
            "type": chunk_type,
            "tokens": len(content.split()),
            "metadata": {
                "chunk_index": chunk_index,
                "timestamp": time.time(),
                "chunk_type": chunk_type,
                **metadata,  # Allow additional metadata to be passed
            },
        }

    def _chunk_text(self, content: str, source_path: str = "") -> list[dict[str, Any]]:
        """Chunk plain text using semantic boundaries and intelligent splitting.

        Implements sophisticated text chunking that respects paragraph boundaries,
        handles variable-length content intelligently, and maintains semantic
        coherence while staying within size limits.

        Chunking Strategies:
            - Semantic: Split by paragraphs, preserve boundaries
            - Size-aware: Handle both character and token limits
            - Overlap: Configurable overlap for context preservation
            - Boundary-aware: Split at natural language boundaries

        Args:
            content: Text content to chunk.
            source_path: Source file path for metadata.

        Returns:
            List of chunk dictionaries with content and metadata.

        Note:
            Uses semantic_chunking configuration to choose between
            paragraph-aware and simple token-based chunking.

        """
        chunks = []

        if self.chunk_config.use_semantic_chunking:
            # Try to split by paragraphs first
            paragraphs = re.split(r"\n\s*\n", content)

            current_chunk = []
            current_size = 0

            for para in paragraphs:
                para_words = para.split()
                len(para_words)
                para_char_size = len(para)

                # If a single paragraph is too large (by characters), split it
                if para_char_size > self.chunk_config.max_chunk_size:
                    # First, save any accumulated chunks
                    if current_chunk:
                        chunks.append(
                            self._create_chunk(
                                content="\n\n".join(current_chunk),
                                chunk_type="semantic",
                                chunk_index=len(chunks),
                                source=source_path,
                                paragraph_count=len(current_chunk),
                            )
                        )
                        current_chunk = []
                        current_size = 0

                    # Split the large paragraph into smaller chunks by characters
                    para_text = para
                    start = 0
                    while start < len(para_text):
                        end = start + self.chunk_config.max_chunk_size
                        if end >= len(para_text):
                            chunk_content = para_text[start:]
                        else:
                            # Try to break at word boundary
                            chunk_content = para_text[start:end]
                            last_space = chunk_content.rfind(" ")
                            if (
                                last_space > 0
                                and last_space > start + self.chunk_config.max_chunk_size // 2
                            ):
                                chunk_content = para_text[start : start + last_space]
                                end = start + last_space

                        if chunk_content.strip():
                            chunks.append(
                                self._create_chunk(
                                    content=chunk_content.strip(),
                                    chunk_type="semantic",
                                    chunk_index=len(chunks),
                                    source=source_path,
                                    paragraph_count=1,
                                    is_split=True,
                                )
                            )

                        # Always advance start to avoid infinite loop
                        start = (
                            end - self.chunk_config.chunk_overlap if end < len(para_text) else end
                        )

                        # Ensure we make progress even if overlap is large
                        if start <= end - self.chunk_config.max_chunk_size:
                            start = end
                elif current_size + para_char_size > self.chunk_config.max_chunk_size:
                    if current_chunk:
                        chunks.append(
                            self._create_chunk(
                                content="\n\n".join(current_chunk),
                                chunk_type="semantic",
                                chunk_index=len(chunks),
                                source=source_path,
                                paragraph_count=len(current_chunk),
                            )
                        )
                    current_chunk = [para]
                    current_size = para_char_size
                else:
                    current_chunk.append(para)
                    current_size += para_char_size

            # Add final chunk
            if current_chunk:
                final_content = "\n\n".join(current_chunk)
                # If final chunk is still too large, split it
                if len(final_content) > self.chunk_config.max_chunk_size:
                    start = 0
                    while start < len(final_content):
                        end = start + self.chunk_config.max_chunk_size
                        if end >= len(final_content):
                            chunk_content = final_content[start:]
                        else:
                            chunk_content = final_content[start:end]
                            last_space = chunk_content.rfind(" ")
                            if last_space > 0:
                                chunk_content = final_content[start : start + last_space]
                                end = start + last_space

                        if chunk_content.strip():
                            chunks.append(
                                self._create_chunk(
                                    content=chunk_content.strip(),
                                    chunk_type="semantic",
                                    chunk_index=len(chunks),
                                    source=source_path,
                                    paragraph_count=len(current_chunk),
                                    is_split=True,
                                )
                            )

                        start = (
                            end - self.chunk_config.chunk_overlap
                            if end < len(final_content)
                            else end
                        )
                else:
                    chunks.append(
                        self._create_chunk(
                            content=final_content,
                            chunk_type="semantic",
                            chunk_index=len(chunks),
                            source=source_path,
                            paragraph_count=len(current_chunk),
                        )
                    )
        else:
            # Simple token-based chunking
            words = content.split()
            chunk_size = self.chunk_config.max_chunk_size
            overlap = self.chunk_config.chunk_overlap

            for i in range(0, len(words), chunk_size - overlap):
                chunk_words = words[i : i + chunk_size]
                if chunk_words:
                    chunks.append(
                        self._create_chunk(
                            content=" ".join(chunk_words),
                            chunk_type="token",
                            chunk_index=len(chunks),
                            source=source_path,
                            word_count=len(chunk_words),
                        )
                    )

        return chunks

    async def process_directory(self, directory: Path) -> list[ProcessedDocument]:
        """Process all matching files in a directory tree with parallel processing.

        Recursively processes all files in the directory that match configured
        file patterns, applying appropriate processing strategies for each file type.
        Handles errors gracefully and continues processing remaining files.

        Args:
            directory: Root directory path to process. Must exist and be readable.

        Returns:
            List of ProcessedDocument objects for all successfully processed files.
            Failed files are logged but don't stop processing of other files.

        Example:
            Process an entire project directory:

            >>> from pathlib import Path
            >>>
            >>> # Process project with mixed file types
            >>> documents = await processor.process_directory(Path("/project"))
            >>>
            >>> # Analyze results by type
            >>> by_type = {}
            >>> for doc in documents:
            ...     doc_type = doc.doc_type
            ...     by_type[doc_type] = by_type.get(doc_type, 0) + 1
            >>>
            >>> print(f"Processed {len(documents)} files:")
            >>> for doc_type, count in by_type.items():
            ...     print(f"  {doc_type}: {count} files")
            >>>
            >>> # Find largest documents
            >>> sorted_docs = sorted(documents, key=lambda d: len(d.content), reverse=True)
            >>> print(f"Largest: {sorted_docs[0].file_path} ({len(sorted_docs[0].content)} chars)")

        Note:
            Processing continues even if individual files fail. Check logs for
            specific file processing errors.

        """
        documents = []

        for pattern in self.doc_config.file_patterns:
            for file_path in directory.rglob(pattern):
                if file_path.is_file():
                    try:
                        doc = await self.process_file(file_path)
                        if doc:
                            documents.append(doc)
                            logger.info(f"Processed {file_path}")
                    except Exception as e:
                        logger.error(f"Error processing {file_path}: {e}")

        return documents
